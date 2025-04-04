from flask import Flask, request, jsonify, render_template, redirect, url_for, session, flash, make_response
import requests
import json
import sqlite3
import os
import re
from functools import wraps
from werkzeug.security import generate_password_hash, check_password_hash
import uuid
from datetime import datetime
import textstat
import nltk
from nltk.sentiment import SentimentIntensityAnalyzer
from docx import Document
import pdfkit
from io import BytesIO
import base64
from docx2pdf import convert
import tempfile


# Download NLTK resources
try:
    nltk.data.find('vader_lexicon')
except LookupError:
    nltk.download('vader_lexicon')

from flask_wtf.csrf import CSRFProtect

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'your_default_secret_key')
csrf = CSRFProtect(app)

# LM Studio API configuration
LM_STUDIO_API_URL = "http://localhost:1234/v1/chat/completions"

# Database setup
def get_db_connection():
    conn = sqlite3.connect('resume_builder.db')
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Users table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        email TEXT UNIQUE NOT NULL,
        password TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    
    # Resumes table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS resumes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        content TEXT NOT NULL,
        score REAL,
        target_job TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users (id)
    )
    ''')
    
    # Resume sections table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS resume_sections (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        resume_id INTEGER NOT NULL,
        section_name TEXT NOT NULL,
        content TEXT NOT NULL,
        FOREIGN KEY (resume_id) REFERENCES resumes (id) ON DELETE CASCADE
    )
    ''')
    
    # Templates table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS templates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        description TEXT,
        css_content TEXT,
        html_structure TEXT
    )
    ''')
    
    # Insert some default templates
    cursor.execute('''
    INSERT OR IGNORE INTO templates (id, name, description, css_content, html_structure)
    VALUES 
    (1, 'Professional', 'Clean and professional template suitable for most industries', 
    'body { font-family: Arial, sans-serif; } .section { margin-bottom: 20px; } h1 { color: #2c3e50; } h2 { color: #3498db; border-bottom: 1px solid #e0e0e0; }',
    '<div class="resume"><h1>{name}</h1><p>{contact}</p>{sections}</div>'),
    
    (2, 'Modern', 'Contemporary layout with accent colors', 
    'body { font-family: "Segoe UI", sans-serif; } .section { margin-bottom: 25px; } h1 { color: #16a085; } h2 { color: #2980b9; padding-bottom: 5px; }',
    '<div class="resume modern"><div class="header"><h1>{name}</h1><p>{contact}</p></div>{sections}</div>'),
    
    (3, 'Minimalist', 'Simple and elegant with focus on content', 
    'body { font-family: Georgia, serif; line-height: 1.6; } .section { margin-bottom: 18px; } h1 { font-weight: normal; } h2 { font-weight: normal; color: #555; }',
    '<div class="resume minimalist"><h1>{name}</h1><p>{contact}</p>{sections}</div>')
    ''')
    
    # Conversations table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS conversations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        resume_id INTEGER,
        messages TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users (id),
        FOREIGN KEY (resume_id) REFERENCES resumes (id)
    )
    ''')
    
    conn.commit()
    conn.close()

# Initialize database on startup
init_db()

# Authentication decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login', next=request.url))
        return f(*args, **kwargs)
    return decorated_function

# Create chat completion using LM Studio API
def create_chat_completion(user_input, context=""):
    """
    Creates a chat completion using the LM Studio API
    With optional context for more personalized responses
    """
    # For resume-specific guidance with additional context
    resume_instructions = f"""
    You are a helpful resume-building assistant. Help the user create or improve their resume
    by providing specific, actionable advice on formatting, content, and wording.
    Be encouraging but honest. Focus on personalized recommendations.
    Use strong action verbs and help quantify achievements. 
    
    {context}
    
    User query: 
    """
    
    enhanced_input = resume_instructions + user_input
    
    headers = {
        "Content-Type": "application/json"
    }
    data = {
        "model": "mistral-7b-instruct-v0.3:2",
        "messages": [
            {"role": "user", "content": enhanced_input}
        ],
        "temperature": 0.7,
        "max_tokens": -1,
        "stream": False
    }
    
    response = requests.post(LM_STUDIO_API_URL, headers=headers, data=json.dumps(data))

    if response.status_code == 200:
        return response.json()
    else:
        raise Exception(f"Error {response.status_code}: {response.text}")

# ATS Keyword Optimization
def optimize_for_ats(resume_text, job_description):
    """Analyze and optimize a resume for ATS compatibility based on job description"""
    prompt = f"""
    Analyze this resume for ATS (Applicant Tracking System) compatibility based on the job description.
    Identify missing keywords and suggest improvements.
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    
    Please provide:
    1. A list of important keywords from the job description that are missing in the resume
    2. Specific suggestions for incorporating these keywords naturally
    3. Overall ATS optimization score out of 100
    """
    
    try:
        completion = create_chat_completion(prompt)
        return completion['choices'][0]['message']['content']
    except Exception as e:
        return f"Error in ATS analysis: {str(e)}"

# Grammar and formatting check
def check_grammar_formatting(text):
    """Perform basic grammar and formatting checks on resume text"""
    prompt = f"""
    Analyze this resume text for grammatical errors, formatting inconsistencies, and style improvements:
    
    {text}
    
    Please provide:
    1. Grammatical errors with corrections
    2. Formatting inconsistencies and how to fix them
    3. Style improvements for stronger impact
    """
    
    try:
        completion = create_chat_completion(prompt)
        return completion['choices'][0]['message']['content']
    except Exception as e:
        return f"Error in grammar check: {str(e)}"

# Resume scoring system
def score_resume(resume_text, job_title=None):
    """Score a resume based on various factors and return score with feedback"""
    # Basic metrics
    word_count = len(resume_text.split())
    
    # Readability
    readability_score = textstat.flesch_reading_ease(resume_text)
    
    # Action verbs list (simplified)
    action_verbs = [
        'achieved', 'improved', 'launched', 'developed', 'created', 
        'implemented', 'managed', 'led', 'designed', 'increased',
        'decreased', 'reduced', 'negotiated', 'coordinated', 'generated',
        'delivered', 'organized', 'supervised', 'trained', 'analyzed'
    ]
    
    # Count action verbs
    action_verb_count = sum(1 for word in resume_text.lower().split() if word in action_verbs)
    
    # Count bullet points
    bullet_point_count = resume_text.count('•') + resume_text.count('-') + resume_text.count('*')
    
    # Job-specific scoring if job title is provided
    job_specific_score = 0
    if job_title:
        prompt = f"""
        Score this resume for the position of {job_title} on a scale of 0-100.
        Consider relevance, qualifications, and presentation.
        
        Resume:
        {resume_text}
        
        Provide a numerical score and very brief explanation.
        """
        try:
            completion = create_chat_completion(prompt)
            response = completion['choices'][0]['message']['content']
            # Extract score from response (simple approach)
            match = re.search(r'(\d{1,3})(?:/100)?', response)
            if match:
                job_specific_score = int(match.group(1))
            else:
                job_specific_score = 60  # Default if parsing fails
        except:
            job_specific_score = 60  # Default on error
    
    # Calculate final score (weighted components)
    base_score = 50
    length_score = min(20, max(0, (word_count - 200) / 20)) if word_count < 600 else max(0, 20 - (word_count - 600) / 50)
    readability_bonus = min(10, max(0, (readability_score - 30) / 5))
    action_verb_bonus = min(10, action_verb_count / 2)
    bullet_point_bonus = min(10, bullet_point_count / 3)
    
    total_score = base_score + length_score + readability_bonus + action_verb_bonus + bullet_point_bonus
    
    if job_title:
        total_score = (total_score + job_specific_score) / 2
    
    # Round to integer
    total_score = round(min(100, max(0, total_score)))
    
    # Generate feedback
    feedback = {
        'score': total_score,
        'length': 'Good length' if 300 <= word_count <= 700 else ('Too short' if word_count < 300 else 'Too long'),
        'readability': 'Easy to read' if readability_score > 50 else 'Could be more readable',
        'action_verbs': f'Used {action_verb_count} action verbs' + (' (good)' if action_verb_count >= 10 else ' (needs more)'),
        'bullet_points': f'Contains {bullet_point_count} bullet points' + (' (good)' if bullet_point_count >= 15 else ' (consider adding more)'),
    }
    
    return feedback

# Basic sentiment analysis
def analyze_sentiment(text):
    """Perform basic sentiment analysis on resume text"""
    sia = SentimentIntensityAnalyzer()
    sentiment_scores = sia.polarity_scores(text)
    
    # Interpret scores
    if sentiment_scores['compound'] >= 0.05:
        tone = "positive"
    elif sentiment_scores['compound'] <= -0.05:
        tone = "negative"
    else:
        tone = "neutral"
    
    # Resume-specific advice based on sentiment
    if tone == "positive":
        advice = "Your resume has a positive tone, which is generally good. Make sure achievements are backed by specific metrics."
    elif tone == "negative":
        advice = "Your resume has a somewhat negative tone. Try to frame challenges as opportunities and use more positive language."
    else:
        advice = "Your resume has a neutral tone. Consider adding more dynamic language and achievement-focused content."
    
    return {
        'tone': tone,
        'scores': sentiment_scores,
        'advice': advice
    }

# Routes
@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        
        # Hash password
        hashed_password = generate_password_hash(password)
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Check if username or email already exists
        cursor.execute('SELECT * FROM users WHERE username = ? OR email = ?', (username, email))
        if cursor.fetchone():
            conn.close()
            flash('Username or email already exists. Please choose another or login.')
            return redirect(url_for('register'))
        
        # Insert new user
        cursor.execute(
            'INSERT INTO users (username, email, password) VALUES (?, ?, ?)',
            (username, email, hashed_password)
        )
        
        conn.commit()
        user_id = cursor.lastrowid
        conn.close()
        
        # Set session
        session['user_id'] = user_id
        session['username'] = username
        
        flash('Registration successful!')
        return redirect(url_for('dashboard'))
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM users WHERE username = ?', (username,))
        user = cursor.fetchone()
        conn.close()
        
        if user and check_password_hash(user['password'], password):
            session['user_id'] = user['id']
            session['username'] = user['username']
            
            next_page = request.args.get('next')
            if next_page:
                return redirect(next_page)
            
            return redirect(url_for('dashboard'))
        
        flash('Invalid username or password')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out')
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard():
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get user's resumes
    cursor.execute('SELECT * FROM resumes WHERE user_id = ? ORDER BY updated_at DESC', (session['user_id'],))
    resumes = cursor.fetchall()
    
    conn.close()
    
    return render_template('dashboard.html', resumes=resumes)

@app.route('/resume/new', methods=['GET', 'POST'])
@login_required
def new_resume():
    if request.method == 'POST':
        title = request.form['title']
        target_job = request.form['target_job'] if 'target_job' in request.form else ''
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Create new resume
        cursor.execute(
            'INSERT INTO resumes (user_id, title, content, target_job) VALUES (?, ?, ?, ?)',
            (session['user_id'], title, '', target_job)
        )
        
        resume_id = cursor.lastrowid
        
        # Initialize default sections
        default_sections = [
            ('Personal Information', ''),
            ('Work Experience', ''),
            ('Education', ''),
            ('Skills', ''),
            ('Additional Sections', '')
        ]
        
        for section_name, content in default_sections:
            cursor.execute(
                'INSERT INTO resume_sections (resume_id, section_name, content) VALUES (?, ?, ?)',
                (resume_id, section_name, content)
            )
        
        conn.commit()
        conn.close()
        
        return redirect(url_for('edit_resume', resume_id=resume_id))
    
    return render_template('new_resume.html')

@app.route('/resume/<int:resume_id>')
@login_required
def view_resume(resume_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get resume
    cursor.execute('SELECT * FROM resumes WHERE id = ? AND user_id = ?', (resume_id, session['user_id']))
    resume = cursor.fetchone()
    
    if not resume:
        conn.close()
        flash('Resume not found or access denied')
        return redirect(url_for('dashboard'))
    
    # Get sections
    cursor.execute('SELECT * FROM resume_sections WHERE resume_id = ?', (resume_id,))
    sections = cursor.fetchall()
    
    # Get templates
    cursor.execute('SELECT id, name, description FROM templates')
    templates = cursor.fetchall()
    
    conn.close()
    
    return render_template('view_resume.html', resume=resume, sections=sections, templates=templates)

@app.route('/resume/<int:resume_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_resume(resume_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get resume
    cursor.execute('SELECT * FROM resumes WHERE id = ? AND user_id = ?', (resume_id, session['user_id']))
    resume = cursor.fetchone()
    
    if not resume:
        conn.close()
        flash('Resume not found or access denied')
        return redirect(url_for('dashboard'))
    
    # Get sections
    cursor.execute('SELECT * FROM resume_sections WHERE resume_id = ?', (resume_id,))
    sections = cursor.fetchall()
    
    if request.method == 'POST':
        # Update resume title and target job
        title = request.form['title']
        target_job = request.form.get('target_job', '')
        
        cursor.execute(
            'UPDATE resumes SET title = ?, target_job = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?', 
            (title, target_job, resume_id)
        )
        
        # Update sections
        for section in sections:
            section_content = request.form.get(f'section_{section["id"]}', '')
            cursor.execute(
                'UPDATE resume_sections SET content = ? WHERE id = ?',
                (section_content, section["id"])
            )
        
        # Compile full content for scoring
        full_content = "\n\n".join(request.form.get(f'section_{section["id"]}', '') for section in sections)
        
        # Score resume
        score_result = score_resume(full_content, target_job)

        # Ensure `score_result` is a dictionary
        if isinstance(score_result, int):
            score_result = {"score": score_result}  # Convert to dictionary
        
        # Update resume content and score
        cursor.execute(
            'UPDATE resumes SET content = ?, score = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?',
            (full_content, score_result['score'], resume_id)
        )
        
        conn.commit()
        conn.close()
        
        flash('Resume updated successfully!')
        return redirect(url_for('view_resume', resume_id=resume_id))
    
    conn.close()
    
    return render_template('edit_resume.html', resume=resume, sections=sections)

@app.route('/resume/<int:resume_id>/chat', methods=['GET', 'POST'])
@login_required
def resume_chat(resume_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get resume
    cursor.execute('SELECT * FROM resumes WHERE id = ? AND user_id = ?', (resume_id, session['user_id']))
    resume = cursor.fetchone()
    
    if not resume:
        conn.close()
        flash('Resume not found or access denied')
        return redirect(url_for('dashboard'))
    
    # Get sections
    cursor.execute('SELECT * FROM resume_sections WHERE resume_id = ?', (resume_id,))
    sections = cursor.fetchall()
    
    # Get or initialize conversation
    cursor.execute('SELECT * FROM conversations WHERE resume_id = ? AND user_id = ? ORDER BY created_at DESC LIMIT 1', 
                  (resume_id, session['user_id']))
    conversation = cursor.fetchone()
    
    if conversation:
        messages = json.loads(conversation['messages'])
    else:
        # Start new conversation
        initial_message = {
            "role": "assistant", 
            "content": "Hi! I'm your resume building assistant. I'll help you create or improve your resume. Let's get started! What part of your resume would you like to work on first?"
        }
        messages = [initial_message]
        cursor.execute(
            'INSERT INTO conversations (user_id, resume_id, messages) VALUES (?, ?, ?)',
            (session['user_id'], resume_id, json.dumps(messages))
        )
        conn.commit()
        conversation_id = cursor.lastrowid
    
    if request.method == 'POST':
        user_input = request.form['user_input']
        
        # Add user message
        messages.append({"role": "user", "content": user_input})
        
        # Create context based on resume content
        context = f"User is working on a resume titled '{resume['title']}'"
        if resume['target_job']:
            context += f" for the position of {resume['target_job']}"
        
        # Get AI response
        try:
            completion = create_chat_completion(user_input, context)
            ai_response = completion['choices'][0]['message']['content']
            
            # Add AI message
            messages.append({"role": "assistant", "content": ai_response})
            
            # Update conversation in database
            if conversation:
                cursor.execute(
                    'UPDATE conversations SET messages = ? WHERE id = ?',
                    (json.dumps(messages), conversation['id'])
                )
            else:
                cursor.execute(
                    'INSERT INTO conversations (user_id, resume_id, messages) VALUES (?, ?, ?)',
                    (session['user_id'], resume_id, json.dumps(messages))
                )
            
            conn.commit()
            
            # Check if response contains section updates
            update_section = None
            for section in sections:
                section_name = section['section_name'].lower()
                if section_name in user_input.lower():
                    update_section = section
                    break
            
            # If discussing a specific section, analyze for potential updates
            if update_section and ("improve" in user_input.lower() or "update" in user_input.lower() or "add" in user_input.lower()):
                # Parse AI suggestions into potential content (simplified approach)
                suggested_content = ai_response
                
                # For demo purposes, we'll just make this available to the template
                # In a real implementation, you might use NLP to extract structured suggestions
                pass
            
        except Exception as e:
            flash(f'Error getting response: {str(e)}')
    
    conn.close()
    
    return render_template('resume_chat.html', resume=resume, sections=sections, messages=messages)

@app.route('/resume/<int:resume_id>/analyze', methods=['GET', 'POST'])
@login_required
def analyze_resume(resume_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get resume
    cursor.execute('SELECT * FROM resumes WHERE id = ? AND user_id = ?', (resume_id, session['user_id']))
    resume = cursor.fetchone()
    
    if not resume:
        conn.close()
        flash('Resume not found or access denied')
        return redirect(url_for('dashboard'))
    
    # Get sections
    cursor.execute('SELECT * FROM resume_sections WHERE resume_id = ?', (resume_id,))
    sections = cursor.fetchall()
    
    analysis = {}
    
    if request.method == 'POST' or resume['content']:
        content = resume['content']
        job_description = request.form.get('job_description', '')
        
        # Basic analyses
        analysis['score'] = score_resume(content, resume['target_job'])
        analysis['length'] = f"{len(content.split())} words"
        analysis['readability'] = check_readability(content)
        analysis['action_verbs'] = f"Used {count_action_verbs(content)} action verbs"
        analysis['bullet_points'] = f"Contains {content.count('•')} bullet points"
        
        # Additional analyses
        analysis['strengths'] = identify_strengths(content, resume['target_job'])
        analysis['areas_for_improvement'] = identify_improvements(content, resume['target_job'])
        analysis['section_feedback'] = generate_section_feedback(sections)
        analysis['recommendations'] = generate_recommendations(content, resume['target_job'])
        
        if job_description:
            analysis['ats'] = optimize_for_ats(content, job_description)
        
        analysis['grammar'] = check_grammar_formatting(content)
    
    conn.close()
    
    return render_template('analyze_resume.html', resume=resume, sections=sections, analysis=analysis)

# Helper functions
def check_readability(content):
    # Simple readability check (can be enhanced with proper readability algorithms)
    words = content.split()
    if len(words) < 200:
        return "Too short"
    elif len(words) > 800:
        return "Too long"
    else:
        return "Could be more readable"

def count_action_verbs(content):
    # Simple action verb counter (can be enhanced with a comprehensive list)
    action_verbs = ["developed", "created", "managed", "led", "implemented", "designed", 
                   "coordinated", "achieved", "improved", "reduced", "increased", "generated"]
    count = 0
    for verb in action_verbs:
        count += content.lower().count(f" {verb} ")
    return count

def identify_strengths(content, target_job):
    # Identify strengths based on content and target job
    strengths = []
    
    # Technical skills check
    tech_skills = ["python", "java", "javascript", "react", "node", "sql", "aws", "docker"]
    for skill in tech_skills:
        if skill in content.lower() and skill in target_job.lower():
            strengths.append(f"Matching {skill.upper()} skills")
    
    # Experience check
    experience_indicators = ["years of experience", "led team", "managed project"]
    for indicator in experience_indicators:
        if indicator in content.lower():
            strengths.append("Demonstrated leadership experience")
            break
    
    # If no strengths found
    if not strengths:
        strengths.append("Add more relevant skills and experience to highlight strengths")
    
    return strengths

def identify_improvements(content, target_job):
    # Identify areas for improvement
    improvements = []
    
    # Length check
    word_count = len(content.split())
    if word_count < 300:
        improvements.append("Resume is too short - add more relevant details")
    elif word_count > 700:
        improvements.append("Resume is too long - consider condensing")
    
    # Action verbs check
    if count_action_verbs(content) < 8:
        improvements.append("Use more action verbs to strengthen impact")
    
    # Keywords check
    if target_job:
        job_keywords = extract_keywords(target_job)
        missing_keywords = [k for k in job_keywords if k not in content.lower()]
        if missing_keywords:
            improvements.append(f"Missing key terms: {', '.join(missing_keywords[:3])}")
    
    # If no improvements found
    if not improvements:
        improvements.append("Continue refining your resume for optimal impact")
    
    return improvements

def extract_keywords(text):
    # Simple keyword extractor (can be enhanced with NLP)
    common_keywords = ["python", "java", "javascript", "agile", "scrum", "leadership", 
                      "development", "software", "engineering", "team", "project"]
    return [word for word in common_keywords if word in text.lower()]

def generate_section_feedback(sections):
    feedback = {}
    for section in sections:
        section_type = section['type'].lower() if 'type' in section else ''
        section_content = section['content'] if 'content' in section else ''
        
        if section_type == 'experience':
            feedback[section['id']] = {
                'title': 'Experience Section',
                'feedback': analyze_experience_section(section_content)
            }
        elif section_type == 'education':
            feedback[section['id']] = {
                'title': 'Education Section',
                'feedback': analyze_education_section(section_content)
            }
        elif section_type == 'skills':
            feedback[section['id']] = {
                'title': 'Skills Section',
                'feedback': analyze_skills_section(section_content)
            }
        else:
            feedback[section['id']] = {
                'title': section['title'] if 'title' in section else 'Section',
                'feedback': "Consider organizing this section with bullet points for better readability"
            }
    
    return feedback

def analyze_experience_section(content):
    # Analyze experience section
    feedback = []
    
    if len(content.split()) < 100:
        feedback.append("Add more details about your achievements and responsibilities")
    
    if content.count('•') < 3:
        feedback.append("Use bullet points to highlight key achievements")
    
    if not any(month in content for month in ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]):
        feedback.append("Include specific dates for each position")
    
    if not feedback:
        feedback.append("Good job providing detailed experience information")
    
    return feedback

def analyze_education_section(content):
    # Analyze education section
    feedback = []
    
    if "GPA" not in content:
        feedback.append("Consider adding your GPA if it's above 3.0")
    
    if not any(year in content for year in ["2018", "2019", "2020", "2021", "2022", "2023", "2024", "2025"]):
        feedback.append("Include graduation years")
    
    if not feedback:
        feedback.append("Education section looks comprehensive")
    
    return feedback

def analyze_skills_section(content):
    # Analyze skills section
    feedback = []
    
    if len(content.split(',')) < 5 and len(content.split('•')) < 5:
        feedback.append("List more relevant skills (aim for 8-12 key skills)")
    
    if not any(tech in content.lower() for tech in ["python", "java", "javascript", "react", "node", "sql"]):
        feedback.append("Include more technical skills relevant to software engineering")
    
    if not feedback:
        feedback.append("Skills section effectively highlights your capabilities")
    
    return feedback

def generate_recommendations(content, target_job):
    # Generate recommendations
    recommendations = []
    
    # Length recommendations
    word_count = len(content.split())
    if word_count < 350:
        recommendations.append("Expand your resume with more detailed achievements and responsibilities")
    elif word_count > 700:
        recommendations.append("Consider condensing your resume to focus on the most relevant experience")
    
    # ATS optimization
    recommendations.append("Customize your resume keywords to match the specific job description")
    
    # Action verbs
    if count_action_verbs(content) < 10:
        recommendations.append("Strengthen impact by using more action verbs like 'developed', 'implemented', and 'achieved'")
    
    # Quantifiable achievements
    if not any(char.isdigit() for char in content):
        recommendations.append("Add quantifiable achievements (e.g., 'increased efficiency by 20%')")
    
    # Missing sections check
    if "education" not in content.lower():
        recommendations.append("Add an Education section with your degrees and relevant coursework")
    
    if "skills" not in content.lower():
        recommendations.append("Include a dedicated Skills section highlighting technical and soft skills")
    
    return recommendations

def score_resume(content, target_job):
    # Calculate an overall score
    score = 61  # Base score
    
    # Length adjustment
    word_count = len(content.split())
    if 350 <= word_count <= 700:
        score += 10
    elif word_count < 200 or word_count > 1000:
        score -= 10
    
    # Action verbs adjustment
    action_verb_count = count_action_verbs(content)
    if action_verb_count >= 10:
        score += 5
    elif action_verb_count <= 5:
        score -= 5
    
    # Keywords adjustment
    if target_job:
        keywords = extract_keywords(target_job)
        matching_keywords = sum(1 for k in keywords if k in content.lower())
        score += min(matching_keywords * 3, 15)
    
    # Quantifiable achievements
    if sum(c.isdigit() for c in content) > 5:
        score += 5
    
    # Bullet points
    if content.count('•') >= 10:
        score += 5
    elif content.count('•') <= 3:
        score -= 5
    
    # Cap the score
    return max(0, min(score, 100))

def optimize_for_ats(content, job_description):
    # Simple ATS optimization suggestions
    suggestions = []
    
    # Extract keywords from job description
    job_keywords = extract_keywords(job_description)
    job_keywords.extend([
        "experience", "teamwork", "communication", "leadership", 
        "project management", "software development", "problem-solving"
    ])
    
    # Find missing keywords
    missing_keywords = [k for k in job_keywords if k not in content.lower()]
    
    if missing_keywords:
        suggestions.append(f"Add these keywords to improve ATS match: {', '.join(missing_keywords[:5])}")
    
    suggestions.append("Use standard section headings (Experience, Education, Skills)")
    suggestions.append("Remove graphics, tables, and special characters that ATS might not parse correctly")
    
    return suggestions

def check_grammar_formatting(content):
    # Simple grammar and formatting checks
    issues = []
    
    # Check for common grammar issues
    grammar_checks = {
        "i ": "Capitalize 'I'",
        "  ": "Remove double spaces",
        "!": "Avoid exclamation marks in professional documents",
        "very ": "Use stronger words instead of 'very'"
    }
    
    for pattern, suggestion in grammar_checks.items():
        if pattern in content:
            issues.append(suggestion)
    
    # Check for inconsistent spacing
    if re.search(r'\n\n\n', content):
        issues.append("Avoid excessive blank lines")
    
    # Check for consistent tense
    past_tense = ["developed", "created", "managed", "led"]
    present_tense = ["develop", "create", "manage", "lead"]
    
    past_count = sum(content.lower().count(word) for word in past_tense)
    present_count = sum(content.lower().count(word) for word in present_tense)
    
    if past_count > 0 and present_count > 0:
        issues.append("Use consistent verb tense throughout your resume")
    
    if not issues:
        issues.append("No major grammar or formatting issues detected")
    
    return issues

@app.route('/resume/<int:resume_id>/export', methods=['GET', 'POST'])
@login_required
def export_resume(resume_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Get resume
    cursor.execute('SELECT * FROM resumes WHERE id = ? AND user_id = ?', (resume_id, session['user_id']))
    resume = cursor.fetchone()
    
    if not resume:
        conn.close()
        flash('Resume not found or access denied')
        return redirect(url_for('dashboard'))
    
    # Get sections
    cursor.execute('SELECT * FROM resume_sections WHERE resume_id = ?', (resume_id,))
    sections = cursor.fetchall()
    
    # Get templates
    cursor.execute('SELECT * FROM templates')
    templates = cursor.fetchall()
    
    # Check if this is a GET request with format_type parameter (for iframe direct loading)
    is_preview = False
    format_type = None
    template_id = 1  # Default template
    
    if request.method == 'GET' and 'format_type' in request.args:
        format_type = request.args.get('format_type')
        template_id = request.args.get('template_id', 1)
        is_preview = request.args.get('preview', 'false') == 'true'
    
    if request.method == 'POST' or format_type:
        if not format_type:  # If not set from GET params
            format_type = request.form['format_type']
            template_id = request.form.get('template_id', 1)
            is_preview = request.form.get('preview', 'false') == 'true'
        
        # Get selected template
        cursor.execute('SELECT * FROM templates WHERE id = ?', (template_id,))
        template = cursor.fetchone()
        
        # Prepare content
        content = ""
        for section in sections:
            if section['content'].strip():
                content += f"<div class='section'><h2>{section['section_name']}</h2><div>{section['content']}</div></div>"
        
        # Extract personal info for header
        personal_info = ""
        for section in sections:
            if section['section_name'] == 'Personal Information':
                personal_info = section['content']
                break
        
        # Parse name and contact from personal info (simple implementation)
        name = "John Doe"  # Default
        contact = "email@example.com • (555) 123-4567"  # Default
        
        lines = personal_info.split('\n')
        if lines and lines[0].strip():
            name = lines[0].strip()
        if len(lines) > 1:
            contact = ' • '.join(line.strip() for line in lines[1:3] if line.strip())
        
        # Format HTML based on template
        html_structure = template['html_structure']
        html_content = html_structure.replace('{name}', name).replace('{contact}', contact).replace('{sections}', content)
        
        # Add template CSS
        css = template['css_content']
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{resume['title']}</title>
            <style>
                {css}
                @page {{ size: letter; margin: 0.75in; }}
                body {{ margin: 0; padding: 0; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        if format_type == 'pdf':
            try:
                # Import required modules here to ensure COM initialization happens correctly
                import pythoncom
                
                # Initialize COM for this thread - this fixes the CoInitialize error
                pythoncom.CoInitialize()
                
                try:
                    # Create a new Document
                    doc = Document()
                    
                    # Add name as title
                    doc.add_heading(name, 0)
                    
                    # Add contact info
                    doc.add_paragraph(contact)
                    
                    # Add sections
                    for section in sections:
                        if section['content'].strip():
                            doc.add_heading(section['section_name'], 1)
                            doc.add_paragraph(section['content'])
                    
                    # Create temporary files for the conversion process
                    with tempfile.TemporaryDirectory() as temp_dir:
                        # Create temporary file paths
                        docx_path = os.path.join(temp_dir, f"{resume['title']}.docx")
                        pdf_path = os.path.join(temp_dir, f"{resume['title']}.pdf")
                        
                        # Save the DOCX to the temporary location
                        doc.save(docx_path)
                        
                        # Convert DOCX to PDF
                        convert(docx_path, pdf_path)
                        
                        # Read the PDF file
                        with open(pdf_path, 'rb') as pdf_file:
                            pdf_content = pdf_file.read()
                        
                        # Return the PDF file as a response
                        response = make_response(pdf_content)
                        response.headers['Content-Type'] = 'application/pdf'
                        
                        # Set Content-Disposition based on whether this is a preview or download
                        if is_preview:
                            # For preview, use 'inline' to display in browser
                            response.headers['Content-Disposition'] = f'inline; filename={resume["title"]}.pdf'
                        else:
                            # For download, use 'attachment' to prompt download
                            response.headers['Content-Disposition'] = f'attachment; filename={resume["title"]}.pdf'
                        
                        conn.close()
                        return response
                finally:
                    # Always uninitialize COM when done, even if an exception occurred
                    pythoncom.CoUninitialize()
                        
            except Exception as e:
                flash(f'Error generating PDF: {str(e)}')
                conn.close()
                return redirect(url_for('export_resume', resume_id=resume_id))
                        
        elif format_type == 'docx':
            try:
                # Create a new Document
                doc = Document()
                
                # Add name as title
                doc.add_heading(name, 0)
                
                # Add contact info
                doc.add_paragraph(contact)
                
                # Add sections
                for section in sections:
                    if section['content'].strip():
                        doc.add_heading(section['section_name'], 1)
                        doc.add_paragraph(section['content'])
                
                # Save to BytesIO
                file_stream = BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                
                response = make_response(file_stream.read())
                response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                response.headers['Content-Disposition'] = f'attachment; filename={resume["title"]}.docx'
                
                conn.close()
                return response
            except Exception as e:
                flash(f'Error generating DOCX: {str(e)}')
                conn.close()
                return redirect(url_for('export_resume', resume_id=resume_id))
                
        else:
            # HTML preview
            conn.close()
            return render_template('preview_resume.html', resume=resume, html_content=full_html)
    
    # If we get here, it's a GET request without format_type, so show the export options page
    conn.close()
    return render_template('export_resume.html', resume=resume, templates=templates)
    
@app.route('/api/get_response', methods=['POST'])
@login_required
def api_get_response():
    user_input = request.json.get('prompt')
    resume_id = request.json.get('resume_id')
    
    context = ""
    if resume_id:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM resumes WHERE id = ? AND user_id = ?', (resume_id, session['user_id']))
        resume = cursor.fetchone()
        
        if resume:
            context = f"User is working on a resume titled '{resume['title']}'"
            if resume['target_job']:
                context += f" for the position of {resume['target_job']}"
        
        conn.close()
    
    try:
        completion = create_chat_completion(user_input, context)
        response_text = completion['choices'][0]['message']['content']
        return jsonify({'response': response_text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/check_grammar', methods=['POST'])
@login_required
def api_check_grammar():
    text = request.json.get('text', '')
    
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    result = check_grammar_formatting(text)
    return jsonify({'result': result})

@app.route('/api/optimize_ats', methods=['POST'])
@login_required
def api_optimize_ats():
    resume_text = request.json.get('resume_text', '')
    job_description = request.json.get('job_description', '')

    if not resume_text or not job_description:
        return jsonify({'error': 'Missing required parameters'}), 400

    result = optimize_for_ats(resume_text, job_description)
    return jsonify({'result': result})


@app.route('/api/score_resume', methods=['POST'])
@login_required
def api_score_resume():
    resume_text = request.json.get('resume_text', '')
    job_title = request.json.get('job_title', '')

    if not resume_text:
        return jsonify({'error': 'No resume text provided'}), 400

    result = score_resume(resume_text, job_title)
    return jsonify({'result': result})

@app.route('/api/analyze_sentiment', methods=['POST'])
@login_required
def api_analyze_sentiment():
    text = request.json.get('text', '')

    if not text:
        return jsonify({'error': 'No text provided'}), 400

    result = analyze_sentiment(text)
    return jsonify({'result': result})

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_server_error(e):
    return render_template('500.html'), 500

if __name__ == '__main__':
    app.run(debug=True)
